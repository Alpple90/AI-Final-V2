"""Convert report.md to report.docx using python-docx."""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── styles ────────────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=None, color=None, mono=False):
    if bold:   run.bold   = bold
    if italic: run.italic = italic
    if size:   run.font.size = Pt(size)
    if color:  run.font.color.rgb = RGBColor(*color)
    if mono:   run.font.name = 'Courier New'

def add_heading(text, level):
    p = doc.add_heading(text, level=level)
    return p

def set_table_header_row(table):
    """Make first row bold and light-grey shaded."""
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)

# ── helpers ───────────────────────────────────────────────────────────────────
def add_code_block(text):
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        # light grey background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)

def render_inline(para, text):
    """Add runs to an existing paragraph, handling **bold**, *italic*, `code`."""
    # split on bold (**), italic (*), code (`)
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\\?\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
        elif part == r'\*' or part == '*':
            para.add_run('*')
        else:
            para.add_run(part)

def add_paragraph_with_inline(text, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    render_inline(p, text)
    return p

def parse_table_row(line):
    """Parse a markdown table row into cells."""
    cells = [c.strip() for c in line.strip().strip('|').split('|')]
    return cells

def is_separator_row(line):
    return bool(re.match(r'^\|[\s\-|:]+\|$', line.strip()))

# ── parse & render ─────────────────────────────────────────────────────────────
with open('report.md', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # skip image placeholders
    if re.match(r'!\[.*?\]', line):
        i += 1
        continue

    # horizontal rule
    if re.match(r'^---+\s*$', line):
        doc.add_paragraph()
        i += 1
        continue

    # headings
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text  = m.group(2).strip()
        # strip markdown bold from heading text
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        add_heading(text, level=level)
        i += 1
        continue

    # fenced code block
    if line.strip().startswith('```'):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i].rstrip('\n'))
            i += 1
        i += 1  # skip closing ```
        add_code_block('\n'.join(code_lines))
        continue

    # table
    if line.strip().startswith('|') and i + 1 < len(lines) and is_separator_row(lines[i + 1]):
        # collect all table rows
        header_cells = parse_table_row(line)
        i += 2  # skip header + separator
        data_rows = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            if not is_separator_row(lines[i]):
                data_rows.append(parse_table_row(lines[i].rstrip('\n')))
            i += 1

        num_cols = len(header_cells)
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = 'Table Grid'

        for j, cell_text in enumerate(header_cells):
            cell = table.cell(0, j)
            cell.text = ''
            render_inline(cell.paragraphs[0], cell_text)

        for r, row_cells in enumerate(data_rows):
            for j, cell_text in enumerate(row_cells[:num_cols]):
                cell = table.cell(r + 1, j)
                cell.text = ''
                render_inline(cell.paragraphs[0], cell_text)

        set_table_header_row(table)
        doc.add_paragraph()
        continue

    # bullet list
    m = re.match(r'^(\s*[-*])\s+(.*)', line)
    if m:
        text = m.group(2)
        p = doc.add_paragraph(style='List Bullet')
        render_inline(p, text)
        i += 1
        continue

    # numbered list
    m = re.match(r'^\d+\.\s+(.*)', line)
    if m:
        text = m.group(1)
        p = doc.add_paragraph(style='List Number')
        render_inline(p, text)
        i += 1
        continue

    # italic/bold-only line used as caption (starts with *)
    if line.strip().startswith('*') and line.strip().endswith('*') and not line.strip().startswith('**'):
        caption = line.strip().strip('*')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9.5)
        i += 1
        continue

    # empty line
    if line.strip() == '':
        i += 1
        continue

    # normal paragraph
    p = doc.add_paragraph()
    render_inline(p, line.strip())
    i += 1

doc.save('report.docx')
print("Saved report.docx")
